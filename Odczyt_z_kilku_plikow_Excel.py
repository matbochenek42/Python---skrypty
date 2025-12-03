import openpyxl as xl 
import os

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ktory_komputer = Path(r"Nazwa") 

BASE_PATH = ktory_komputer / r"Ścieżka"
zapis_do_pliku = ktory_komputer / r"Ścieża.txt"
zapis_do_worda = ktory_komputer / r"Ścieżka.docx"

#sprawdzanie poprawności odczytania plików
def link_do_pliku(nazwa_pliku, nazwa_arkusza):
    sciezka_pliku = BASE_PATH / f"{nazwa_pliku} Nazwa.xlsx"
    try:
        wb = xl.load_workbook(sciezka_pliku, data_only=True)
        return wb, wb[nazwa_arkusza]
    except Exception:
        print(f"\n❌ Błąd! Nie można otworzyć pliku")
        raise SystemExit


def funkcja1(nazwa_pliku, start = 5, koniec = 10, kol_nazwa = 4, kol_liczby = 5, nazwa_arkusza = "Nazwa"): 
    
    #przypisanie prawidłowej ścieżki do pliku, nazw pliku i arkusza
    
    wb, arkusz = link_do_pliku(nazwa_pliku, nazwa_arkusza)

    #pętla wewnątrz arkusza odczytująca wartości
    lista_0 = [
        arkusz.cell(row, kol_liczby).value
        for row in range(start, koniec)
    ]
    
    wb.close()
    return lista_0

def stat(nazwa_arkusza, start, kol_nazwa, kol_nar, kol_o): 
    
    #przypisanie prawidłowej ścieżki do pliku, nazw pliku i arkusza
    nazwa_pliku = nazwa_arkusza
    wb, arkusz = link_do_pliku(nazwa_pliku, nazwa_arkusza)

    #pętle wewnątrz arkusza odczytująca wartości 
    nazwa = [
        arkusz.cell(row, kol_nazwa).value
        for row in range(start, 1500)
    ]

    liczba= [
        arkusz.cell(row, kol_nar).value
        for row in range(start, 1500)
    ]
    
    liczba2 = [
        arkusz.cell(row, kol_o).value
        for row in range(start, 1500)
    ]

    wb.close()
    
    #pominięcie pustych kolumn i zapis do listy
    lista_1 = list(zip(nazwa,liczba ,liczba2))
    
    pusta_lista = [
        list(row) for row in lista_1 if row[0] in ("tu odpowiednie nazwy")
    ]

    return pusta_lista   
    
    #da się skrócić tę funkcję (1 lista compr zamiast 3)


def funkcja2(start = 3, koniec = 7, kolumna = 5, nazwa_arkusza = "nazwa"): #dane do tabel z projektami
    sciezka_pliku = r"Nazwa.xlsx"
    
    try:
        wb = xl.load_workbook(sciezka_pliku, data_only=True)
        arkusz = wb[nazwa_arkusza]
    except Exception:
        print(f"\nBłąd! Nie można otworzyć pliku")
        raise SystemExit
    else:
        dane_funkcja2 = [
            arkusz.cell(row, kolumna).value
            for row in range(start, koniec)
        ]

        wb.close()
        return dane_funkcja2

def do_worda(): #zapis pliku do worda wraz z formatowaniem
    document = Document()

    #globalna czcionka i rozmiar
    style = document.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"     
    font.size = Pt(11)  

    #otworzenie pliku
    try:
        with open(zapis_do_pliku, "r", encoding="utf-8") as f:
            lines = [line.rstrip() for line in f.readlines()]

        #"biblioteka" z indeksami w liście
        NAGLOWKI = {1, 8, 15, 22, 27, 28}
        JUSTOWANE = {2, 9, 16, 23, 29, 47}
        POGRUBIENIE = {38}
        PUNKTY = [range(3,7), range(10,14), range(17,21), range(30,34), range(40,43)]


        #pętla dla tekstu
        for i, l in enumerate(lines):

            #pogrubienie i wyśrodkowanie
            if i in NAGLOWKI:
                p = document.add_paragraph()
                run = p.add_run(l)
                run.bold = True #pogrubienie
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER #wyśrodkowanie
        

            # lista punktowa
            elif any(i in _ for _ in PUNKTY):
                document.add_paragraph(l, style="List Bullet")

            #pogrubienie całego tekstu
            elif i in POGRUBIENIE:
                p = document.add_paragraph()
                run = p.add_run(l)
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            elif i in JUSTOWANE:
                p = document.add_paragraph(l)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 

            #zwykły tekst
            else:
                p = document.add_paragraph(l)

        document.save(zapis_do_worda)
    
    except Exception:
        print("\n❌ Błąd! Zły link do pliku txt!")
        raise SystemExit
    else:
        print("✅ Udało się zapisać do pliku word")
    

# Wywołanie funkcji (neutralne nazwy zestawów danych)
zmienna_a = funkcja1("DaneA")
zmienna_b = funkcja1("DaneB")
zmienna_c = funkcja1("DaneC")
zmienna_d = funkcja1("DaneD")
zmienna_e = funkcja1("DaneE", 13, 18, 3, 4)

stat_a = stat("DaneA", 100, 10, 11, 12)
stat_b = stat("DaneB", 99, 16, 17, 18)
stat_c = stat("DaneC", 450, 13, 14, 15)
stat_d = stat("DaneD", 500, 14, 15, 17)

proj = [float(x) * 100 for x in funkcja2()]

# Ostateczna wersja tekstu do zmiennau
tekst_do_plikuu = f"""
Zestaw danych A -
{zmienna_a[1]} wpisów, w ... {zmienna_a[2]} oraz {zmienna_a[3]} ...
... – {stat_a[0][1]} w ... {stat_a[0][2]};
... – {stat_a[1][1]} w ... {stat_a[1][2]};
... – {stat_a[2][1]} w ... {stat_a[2][2]};
W ... – {zmienna_a[4]} szt.

Zestaw danych B – 
...  {zmienna_b[1]} ... {zmienna_b[2]} ... {zmienna_b[3]} ... 
...  – {stat_b[0][1]} ...  {stat_b[0][2]};
...  – {stat_b[1][1]} ... {stat_b[1][2]};
...  – {stat_b[2][1]} ...  {stat_b[2][2]};
...  – {zmienna_b[4]} ... 

Zestaw danych C – 
...  {zmienna_c[1]} ...  {zmienna_c[2]}...  {zmienna_c[3]}... 
... – {stat_c[0][1]} w...  {stat_c[0][2]};
...  – {stat_c[1][1]} w ...  {stat_c[1][2]};
...  – {stat_c[2][1]} w ...  {stat_c[2][2]};
...  – {zmienna_c[4]} ... 

Zestaw danych E – 
...  {zmienna_e[0]} ...  {zmienna_e[1]} ... 
... {zmienna_e[2]} ... {zmienna_e[3]} ... 
"""

mat_do_plikuu = f"""
Zestaw danych D – 
... {zmienna_d[1]} r...  {zmienna_d[2]} ...  {zmienna_d[3]} ... 
... – {stat_d[0][1]}...  {stat_d[0][2]};
...  – {stat_d[1][1]} ... {stat_d[1][2]};
... – {stat_d[2][1]}... {stat_d[2][2]};
... i – {zmienna_d[4]} ... 
"""

pozycja = f"""
...  {proj[3]:.2f}% ... 

...  – ... 
...  – {proj[1]:.2f}% ... 
...  – {proj[2]:.2f}% ... 
"""

podsumowanie = f"""
_... {stat_c[0][1]} ... ({(stat_c[0][1] / 110) * 100:.2f}%),
... {stat_c[1][1]} ...  ({(stat_c[1][1] / 110) * 100:.2f}%),
...  {stat_c[2][1]} ...  ({(stat_c[2][1] / 110) * 100:.2f}%).
"""


try:
    with open(zapis_do_pliku,'w',encoding="utf-8") as f: 
        f.write(f"{tekst_do_plikuu}\n\n\n{mat_do_plikuu}\n\n\n{pozycja}\n\n\n{podsumowanie}")
except Exception:
    print("\n❌ Błąd! Zły link do pliku txt!")
    raise SystemExit
else:
    print("✅ Udało się zapisać do pliku txt")

do_worda()

os.startfile(zapis_do_worda)