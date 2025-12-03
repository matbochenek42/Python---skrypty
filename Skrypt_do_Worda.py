import docx as dx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import os

from pathlib import Path


# ścieżki do plików
KTORY_KOMPUTER = Path(r"Nazwa") 
MAT_RT = KTORY_KOMPUTER / "Nazwa" / "Nazwa_pliku.docx"
PROT_RT = KTORY_KOMPUTER / "Nazwa" / "Nazwa_pliku.docx"

# wczytanie pliku
mat = dx.Document(MAT_RT)
rt = dx.Document()

#globalna czcionka i rozmiar
style = rt.styles["Normal"]
font = style.font
font.name = "Times New Roman"     
font.size = Pt(10) 


# dodanie wierszy do worda (nie uwzgledniając tabel) + wyjustowanie

rt.add_heading('Część 1\n', level=1)

for wiersz in mat.paragraphs:

    p = rt.add_paragraph()
    run = p.add_run(wiersz.text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 

rt.add_heading('Część 2\n', level=1) #rozdzielenie obu części

for tabela in mat.tables:
    for wiersz in tabela.rows:
        
        tekst_wiersza = " ".join(cell.text.strip() for cell in wiersz.cells if cell.text.strip()) # połączenie wszystkich komórek w jednym wierszu
        if tekst_wiersza:  
            p = rt.add_paragraph(tekst_wiersza)
    rt.add_paragraph()  # pusty akapit

rt.save(PROT_RT)
os.startfile(PROT_RT)